import flet as ft
from datetime import datetime, timedelta
import os
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main(page: ft.Page):
    # --- MOBİL EKRAN VE GENEL AYARLAR ---
    page.title = "ŞTGA Hesaplama Sistemi v2.1"
    page.theme_mode = ft.ThemeMode.DARK
    page.window.width = 450   
    page.window.height = 850  
    page.padding = 10
    page.scroll = ft.ScrollMode.AUTO  
    
    # --- ARKA PLAN HAFIZA DEĞİŞKENLERİ ---
    suclar_listesi = []
    gecerli_suclar_filtreli = []
    secili_index = None
    nihai_rapor_metni = ""
    
    tavan_asildi_mi = False
    kesinlesen_gun = 0
    ilk_suc_tarihi_str = ""
    sartlar_olustu_mu = False
    yeni_ceza_miktari_metni = ""
    yeni_ceza_gun = 0

    # =========================================================================
    # YARDIMCI MODÜLLER (MESAJ KUTULARI & TARİH FORMATLAMA)
    # =========================================================================
    def goster_mesaj(baslik, icerik, renk=ft.Colors.BLUE_400):
        dlg = ft.AlertDialog(
            title=ft.Text(baslik, color=renk, weight=ft.FontWeight.BOLD, size=16),
            content=ft.Text(icerik, size=14),
            actions=[ft.TextButton("Tamam", on_click=lambda e: kapat_dialog(dlg))]
        )
        page.overlay.append(dlg)
        dlg.open = True
        page.update()

    def kapat_dialog(dialog):
        dialog.open = False
        page.update()

    def tarih_cevir(tarih_str):
        try:
            tarih_str = tarih_str.replace("/", ".").replace("-", ".")
            return datetime.strptime(tarih_str, "%d.%m.%Y").date()
        except ValueError:
            return None

    def rakam_al(val):
        return int(val) if val and val.isdigit() else 0

    def otomatik_tarih_formatla(e):
        metin = e.control.value.replace(".", "").strip()
        if len(metin) == 8 and metin.isdigit():
            formatli = f"{metin[:2]}.{metin[2:4]}.{metin[4:]}"
            e.control.value = formatli
            page.update()
            
    def asil_tarih_degisti(e):
        otomatik_tarih_formatla(e)
        asil_ceza_kutularini_guncelle()

    def get_lehe_bitis_date():
        return datetime(2020, 3, 30).date() if cmb_lehe_bitis.value == "30.03.2020" else datetime(2016, 7, 1).date()

    # =========================================================================
    # DİNAMİK KUTU KİLİTLEME MOTORLARI
    # =========================================================================
    def asil_ceza_kutularini_guncelle(e=None):
        start_lehe = datetime(2005, 6, 1).date()
        end_lehe = get_lehe_bitis_date()
        asil_tarih = tarih_cevir(txt_asil_tarih.value.strip())
        infaz_orani_var = chk_oran_var.value
        
        acik_olacak = False
        if asil_tarih and start_lehe <= asil_tarih <= end_lehe and infaz_orani_var:
            acik_olacak = True
            
        kapali_renk = ft.Colors.with_opacity(0.1, ft.Colors.GREY_500) if page.theme_mode == ft.ThemeMode.DARK else ft.Colors.with_opacity(0.1, ft.Colors.GREY_300)
        
        for kutu in [txt_asil_yil, txt_asil_ay, txt_asil_gun]:
            if acik_olacak:
                kutu.disabled = False
                kutu.bgcolor = ft.Colors.SURFACE
            else:
                kutu.disabled = True
                kutu.bgcolor = kapali_renk
                kutu.value = ""
        page.update()

    def hapis_sure_kutularini_guncelle(e=None):
        acik_olacak = (cmb_ceza_turu.value == "Hapis")
        kapali_renk = ft.Colors.with_opacity(0.1, ft.Colors.GREY_500) if page.theme_mode == ft.ThemeMode.DARK else ft.Colors.with_opacity(0.1, ft.Colors.GREY_300)
        
        for kutu in [txt_yeni_yil, txt_yeni_ay, txt_yeni_gun]:
            if acik_olacak:
                kutu.disabled = False
                kutu.bgcolor = ft.Colors.SURFACE
            else:
                kutu.disabled = True
                kutu.bgcolor = kapali_renk
                kutu.value = ""
        page.update()

    def chk_oran_degisti(e):
        if chk_oran_var.value:
            asil_tarih = tarih_cevir(txt_asil_tarih.value.strip())
            start_lehe = datetime(2005, 6, 1).date()
            end_lehe = get_lehe_bitis_date()
            if not asil_tarih or not (start_lehe <= asil_tarih <= end_lehe):
                goster_mesaj("Bu Alan Kilitli!", f"1/2 İnfaz Oranı seçeneği, SADECE Asıl Suç Tarihi 01.06.2005 ile {cmb_lehe_bitis.value} aralığında olan dosyalar için aktiftir.", ft.Colors.ORANGE_400)
                chk_oran_var.value = False
        asil_ceza_kutularini_guncelle()

    # =========================================================================
    # LİSTEYE EKLEME / SİLME / AYNI TARİH RADARI MANTIĞI
    # =========================================================================
    def tabloyu_yenile():
        tablo_suclar.rows.clear()
        for idx, s in enumerate(suclar_listesi):
            c_gun = s["toplam_gun"]
            y, a, g = c_gun // 365, (c_gun % 365) // 30, (c_gun % 365) % 30
            sure_metni = f"{y}Y {a}A {g}G" if s["tur"] == "Hapis" else "-"
            
            tablo_suclar.rows.append(
                ft.DataRow(
                    cells=[
                        ft.DataCell(ft.Text(s["tarih_str"])),
                        ft.DataCell(ft.Text(s["kasit"])),
                        ft.DataCell(ft.Text(s["tur"])),
                        ft.DataCell(ft.Text(sure_metni)),
                    ],
                    on_select_change=lambda e, index=idx: listeden_sec(index)
                )
            )
        page.update()

    def sucu_giris_kutularini_temizle():
        txt_yeni_tarih.value = ""
        txt_yeni_yil.value = ""
        txt_yeni_ay.value = ""
        txt_yeni_gun.value = ""
        cmb_kasit.value = "Evet"
        cmb_ceza_turu.value = "Hapis"
        
        # Yeni Flet yapısına göre buton içeriği güncelleme
        btn_sucu_ekle.content = ft.Row([ft.Icon(ft.Icons.ADD), ft.Text("Suçu Listeye Ekle")], alignment=ft.MainAxisAlignment.CENTER)
        btn_sucu_ekle.bgcolor = ft.Colors.GREEN_700
        btn_sucu_sil.disabled = True
        hapis_sure_kutularini_guncelle()

    def listeden_sec(idx):
        nonlocal secili_index
        secili_index = idx
        s = suclar_listesi[idx]
        
        txt_yeni_tarih.value = s["tarih_str"]
        cmb_kasit.value = s["kasit"]
        cmb_ceza_turu.value = s["tur"]
        
        c_gun = s["toplam_gun"]
        txt_yeni_yil.value = str(c_gun // 365) if c_gun > 0 else ""
        txt_yeni_ay.value = str((c_gun % 365) // 30) if c_gun > 0 else ""
        txt_yeni_gun.value = str((c_gun % 365) % 30) if c_gun > 0 else ""
        
        # Yeni Flet yapısına göre buton içeriği güncelleme
        btn_sucu_ekle.content = ft.Row([ft.Icon(ft.Icons.EDIT), ft.Text("Suçu GÜNCELLE")], alignment=ft.MainAxisAlignment.CENTER)
        btn_sucu_ekle.bgcolor = ft.Colors.AMBER_700
        btn_sucu_sil.disabled = False
        hapis_sure_kutularini_guncelle()
        page.update()

    def sucu_ekle_tetikle(e):
        nonlocal secili_index
        tarih_str = txt_yeni_tarih.value.strip()
        gecerli_tarih = tarih_cevir(tarih_str)
        
        if not gecerli_tarih:
            goster_mesaj("Hata", "Lütfen geçerli bir tarih girin! (GG.AA.YYYY)", ft.Colors.RED_400)
            return
            
        tarih_str = gecerli_tarih.strftime("%d.%m.%Y")
        kasit = cmb_kasit.value
        tur = cmb_ceza_turu.value
        yil = rakam_al(txt_yeni_yil.value)
        ay = rakam_al(txt_yeni_ay.value)
        gun = rakam_al(txt_yeni_gun.value)
        toplam_gun = (yil * 365) + (ay * 30) + gun

        if secili_index is not None:
            for i, s in enumerate(suclar_listesi):
                if i != secili_index and s["tarih_str"] == tarih_str: 
                    goster_mesaj("Tarih Çakışması", "Listede aynı tarihli başka bir suç var! Üst üste bindiremezsiniz.", ft.Colors.ORANGE_400)
                    return
            suclar_listesi[secili_index] = {"tarih": gecerli_tarih, "tarih_str": tarih_str, "kasit": kasit, "tur": tur, "toplam_gun": toplam_gun}
            secili_index = None
            sucu_giris_kutularini_temizle()
            tabloyu_yenile()
        else:
            if tur == "Hapis" and toplam_gun == 0:
                goster_mesaj("Uyarı", "Hapis cezası seçildiğinde süre boş veya 0 olamaz!", ft.Colors.ORANGE_400)
                return
                
            ayni_tarihli_index = next((i for i, s in enumerate(suclar_listesi) if s["tarih_str"] == tarih_str), None)
            
            if ayni_tarihli_index is not None:
                def radar_cevap(birlestir):
                    kapat_dialog(radar_dlg)
                    if birlestir:
                        eski_gun = suclar_listesi[ayni_tarihli_index]["toplam_gun"]
                        suclar_listesi[ayni_tarihli_index]["toplam_gun"] = eski_gun + toplam_gun
                        goster_mesaj("Bilgi", "Aynı olaya ait cezalar başarıyla BİRLEŞTİRİLDİ.", ft.Colors.GREEN_400)
                    else:
                        suclar_listesi.append({"tarih": gecerli_tarih, "tarih_str": tarih_str, "kasit": kasit, "tur": tur, "toplam_gun": toplam_gun})
                    sucu_giris_kutularini_temizle()
                    tabloyu_yenile()

                radar_dlg = ft.AlertDialog(
                    title=ft.Text("⚠️ Aynı Tarihli Suç Radarı", color=ft.Colors.ORANGE_400, weight=ft.FontWeight.BOLD),
                    content=ft.Text("Bu tarihte başka bir suç zaten var! Bu eylemlerin içeriği nedir?\n\n• BİRLEŞTİR (Aynı Olay): Süreler toplanır.\n• AYRI EKLE (Farklı Olay): Büyük ceza küçük cezayı yutar."),
                    actions=[
                        ft.TextButton("BİRLEŞTİR (Aynı Olay)", on_click=lambda e: radar_cevap(True)),
                        ft.TextButton("AYRI EKLE (Farklı Olay)", on_click=lambda e: radar_cevap(False))
                    ]
                )
                page.overlay.append(radar_dlg)
                radar_dlg.open = True
                page.update()
            else:
                suclar_listesi.append({"tarih": gecerli_tarih, "tarih_str": tarih_str, "kasit": kasit, "tur": tur, "toplam_gun": toplam_gun})
                sucu_giris_kutularini_temizle()
                tabloyu_yenile()

    def sucu_sil_tetikle(e):
        nonlocal secili_index
        if secili_index is not None:
            suclar_listesi.pop(secili_index)
            secili_index = None
            sucu_giris_kutularini_temizle()
            tabloyu_yenile()

    # =========================================================================
    # HESAPLAMA MOTORU
    # =========================================================================
    def hesapla_tetikle(e):
        nonlocal nihai_rapor_metni, tavan_asildi_mi, kesinlesen_gun, ilk_suc_tarihi_str, sartlar_olustu_mu, yeni_ceza_miktari_metni, yeni_ceza_gun
        
        txt_sonuc.value = ""
        nihai_rapor_metni = ""
        btn_kaydet.disabled = True
        
        asil_tarih = tarih_cevir(txt_asil_tarih.value.strip())
        stt = tarih_cevir(txt_stt.value.strip())
        htt = tarih_cevir(txt_htt.value.strip())
        
        if not all([asil_tarih, stt, htt]):
            goster_mesaj("Eksik Bilgi", "Lütfen Asıl Suç Tarihi, ŞTT ve HTT tarihlerini eksiksiz girin.", ft.Colors.RED_400)
            return
            
        if asil_tarih >= stt or stt >= htt:
            goster_mesaj("Tarih Mantık Hatası", "DİKKAT: Kronolojik sıra (Asıl Suç < ŞTT < HTT) kurallarına aykırı giriş yapıldı!", ft.Colors.RED_400)
            return
            
        limit_tarihi = htt
        ozel_durum_metni = ""
        start_lehe = datetime(2005, 6, 1).date()
        end_lehe = get_lehe_bitis_date()
        
        if start_lehe <= asil_tarih <= end_lehe and chk_oran_var.value:
            a_yil = rakam_al(txt_asil_yil.value)
            a_ay = rakam_al(txt_asil_ay.value)
            a_gun = rakam_al(txt_asil_gun.value)
            if (a_yil + a_ay + a_gun) == 0:
                goster_mesaj("Eksik Asıl Ceza", "1/2 Karma içtima denetim süresi (DSBT) hesabı için 'Asıl Suç Cezası'nı girmelisiniz!", ft.Colors.RED_400)
                return
            asil_ceza_gun = (a_yil * 365) + (a_ay * 30) + a_gun
            denetim_gun = (asil_ceza_gun / 2) / 2
            limit_tarihi = htt - timedelta(days=int(denetim_gun))
            ozel_durum_metni = f"ÖZEL DURUM: Eski yasa (1/2) karma içtima kuralı gereği 'Denetim İndirimi' HTT'den geriye doğru çıkartılarak gerçek Denetim Süresi Bitişi (DSBT) {limit_tarihi.strftime('%d.%m.%Y')} olarak hesaplanmıştır.\n\n"

        if not suclar_listesi:
            goster_mesaj("Boş Liste", "Lütfen denetim sürecinde işlenen kasti suçları ekleyin.", ft.Colors.ORANGE_400)
            return

        gecerli_suclar = []
        gecerli_suclar_filtreli.clear()
        ilk_ihlal_tarihi = htt
        tavan_asildi_mi = False
        
        for suc in suclar_listesi:
            if stt < suc["tarih"] <= limit_tarihi and suc["kasit"] == "Evet" and suc["tur"] == "Hapis":
                if suc["tarih"] < ilk_ihlal_tarihi:
                    ilk_ihlal_tarihi = suc["tarih"]
                iki_kati_gun = suc["toplam_gun"] * 2
                bitis_tarihi = suc["tarih"] + timedelta(days=iki_kati_gun)
                if bitis_tarihi > htt:
                    bitis_tarihi = htt
                    tavan_asildi_mi = True
                gecerli_suclar.append({"bas": suc["tarih"], "bit": bitis_tarihi})
                gecerli_suclar_filtreli.append(suc)

        gecerli_suclar_filtreli.sort(key=lambda x: x["tarih"])
        
        if not gecerli_suclar:
            hata_mesaji = "--- ŞTGA ŞARTLARI OLUŞMADI (HUKUKİ MÜTALAA) ---\n\n"
            for i, suc in enumerate(suclar_listesi):
                hata_mesaji += f"[{i+1}. SUÇ - {suc['tarih'].strftime('%d.%m.%Y')}] NEDEN ELENDİ?\n"
                if not (stt < suc["tarih"] <= limit_tarihi):
                    hata_mesaji += f"- Suç tarihi, yasal denetim süresi dışındadır (Sınır: {limit_tarihi.strftime('%d.%m.%Y')})\n"
                if suc["kasit"] != "Evet":
                    hata_mesaji += "- Suç kasten işlenmemiştir (Taksirlidir).\n"
                if suc["tur"] in ["HAGB", "APC"]:
                    hata_mesaji += f"- Karar türü {suc['tur']} olduğundan mahkumiyet sonucu doğurmaz.\n"
            hata_mesaji += "\nÖZET: Şartları taşıyan ihlal bulunmadığından ŞTGA'ya yer olmadığı mütalaa olunur."
            txt_sonuc.value = hata_mesaji
            nihai_rapor_metni = hata_mesaji
            sartlar_olustu_mu = False
            btn_kaydet.disabled = False
            page.update()
            return

        gecerli_suclar.sort(key=lambda x: x["bas"])
        coklu_suc_var = len(gecerli_suclar) > 1
        kesisme_var = False
        
        total_stga_gun = 0
        mevcut_bas = gecerli_suclar[0]["bas"]
        mevcut_bit = gecerli_suclar[0]["bit"]
        
        for i in range(1, len(gecerli_suclar)):
            suc = gecerli_suclar[i]
            if suc["bas"] <= mevcut_bit:
                kesisme_var = True
                if suc["bit"] > mevcut_bit:
                    mevcut_bit = suc["bit"]
            else:
                total_stga_gun += (mevcut_bit - mevcut_bas).days
                mevcut_bas = suc["bas"]
                mevcut_bit = suc["bit"]
                
        total_stga_gun += (mevcut_bit - mevcut_bas).days
        genel_tavan_gun = (htt - ilk_ihlal_tarihi).days
        
        if total_stga_gun >= genel_tavan_gun:
            total_stga_gun = genel_tavan_gun
            tavan_asildi_mi = True

        kesinlesen_gun = total_stga_gun
        ilk_suc_tarihi_str = ilk_ihlal_tarihi.strftime("%d.%m.%Y")
        sartlar_olustu_mu = True
        
        ilk_suc_obj = next((s for s in suclar_listesi if s["tarih"] == ilk_ihlal_tarihi), None)
        if ilk_suc_obj:
            cg = ilk_suc_obj["toplam_gun"]
            yeni_ceza_gun = cg
            c_y, c_a, c_g = cg // 365, (cg % 365) // 30, (cg % 365) % 30
            parcalar = []
            if c_y > 0: parcalar.append(f"{c_y} YIL")
            if c_a > 0: parcalar.append(f"{c_a} AY")
            if c_g > 0: parcalar.append(f"{c_g} GÜN")
            yeni_ceza_miktari_metni = " ".join(parcalar) if parcalar else "0 GÜN"

        s_yil = total_stga_gun // 365
        kalan = total_stga_gun % 365
        s_ay = kalan // 30
        s_gun = kalan % 30
        
        rapor = "--- ŞTGA HESAPLAMA RAPORU ---\n\n"
        if ozel_durum_metni:
            rapor += ozel_durum_metni
            
        if not coklu_suc_var:
            aciklama = "AÇIKLAMA: Hükümlünün denetim süresi içerisinde kasten işlemiş olduğu yeni suç incelenmiş olup, yasa gereği bu cezanın 2 katı hesaplanmıştır."
        else:
            if not kesisme_var:
                aciklama = "AÇIKLAMA: Listedeki suçlar kronolojik sıraya dizilmiş, ardışık cezaların infaz taşımaları toplanarak boşluklar atlanmıştır (Kronolojik İçtima)."
            else:
                aciklama = "AÇIKLAMA: Suçlar tarih sırasına dizilmiş, çakışan infaz sürelerinde büyük cezanın küçüğü yutması (Kesişme Kuralı) işletilmiştir."
                
        if tavan_asildi_mi:
            aciklama += " ANCAK hesaplanan bu sürenin Bihakkın Tahliye Tarihini (HTT) aşması nedeniyle; sınır olan 'İlk Suçtan HTT'ye kadar kalan süre' tavan olarak alınmıştır."
            
        rapor += aciklama + "\n\n"
        rapor += f"KESİNLEŞEN İNFAZ MİKTARI:\n{total_stga_gun} GÜN\n(Takribi: {s_yil} Yıl, {s_ay} Ay, {s_gun} Gün)"
        
        txt_sonuc.value = rapor
        nihai_rapor_metni = rapor
        btn_kaydet.disabled = False
        page.update()

    def ekrani_temizle_tetikle(e):
        def temizlik_onay(evet):
            kapat_dialog(temizle_dlg)
            if evet:
                txt_asil_tarih.value = ""
                txt_stt.value = ""
                txt_htt.value = ""
                txt_asil_yil.value = ""
                txt_asil_ay.value = ""
                txt_asil_gun.value = ""
                chk_oran_var.value = False
                suclar_listesi.clear()
                gecerli_suclar_filtreli.clear()
                txt_sonuc.value = ""
                btn_kaydet.disabled = True
                sucu_giris_kutularini_temizle()
                tabloyu_yenile()
                asil_ceza_kutularini_guncelle()

        temizle_dlg = ft.AlertDialog(
            title=ft.Text("🧹 Temizleme Onayı"),
            content=ft.Text("Ekrandaki tüm veriler ve suç listesi tamamen silinecektir. Emin misiniz?"),
            actions=[
                ft.TextButton("Evet", on_click=lambda e: temizlik_onay(True)),
                ft.TextButton("Hayır", on_click=lambda e: temizlik_onay(False))
            ]
        )
        page.overlay.append(temizle_dlg)
        temizle_dlg.open = True
        page.update()

    # =========================================================================
    # BİLGİ PENCERELERİ (REHBER & HAKKINDA) VE DOSYA İŞLEMLERİ
    # =========================================================================
    def evrak_secim_ekrani(e):
        def mutalaa_sec(e):
            kapat_dialog(secim_dlg)
            try:
                doc = Document()
                baslik = doc.add_paragraph()
                baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
                baslik.add_run("ŞTGA HESAPLAMA VE DEĞERLENDİRME MÜTALAASI").bold = True
                p = doc.add_paragraph()
                p.add_run(nihai_rapor_metni)
                doc.save("STGA_Mobil_Mutalaasi.docx")
                goster_mesaj("Başarılı", "Mütalaa, program klasörüne 'STGA_Mobil_Mutalaasi.docx' adıyla kaydedildi.", ft.Colors.GREEN_400)
            except Exception as ex:
                goster_mesaj("Hata", f"Belge oluşturulamadı: {ex}", ft.Colors.RED_400)

        secim_dlg = ft.AlertDialog(
            title=ft.Text("📄 Evrak Çıktısı Seçimi"),
            content=ft.Text("Hangi evrakı oluşturmak istiyorsunuz?"),
            actions=[
                ft.ElevatedButton("1- Gerekçeli ŞTGA Mütalaası", on_click=mutalaa_sec, bgcolor=ft.Colors.GREEN_700, color=ft.Colors.WHITE),
                ft.ElevatedButton("Kapat", on_click=lambda e: kapat_dialog(secim_dlg))
            ]
        )
        page.overlay.append(secim_dlg)
        secim_dlg.open = True
        page.update()

    def taslak_disa_aktar(e):
        veri = {
            "asil_tarih": txt_asil_tarih.value.strip(),
            "lehe_bitis": cmb_lehe_bitis.value,
            "oran_var": 1 if chk_oran_var.value else 0,
            "stt": txt_stt.value.strip(),
            "htt": txt_htt.value.strip(),
            "asil_yil": txt_asil_yil.value.strip(),
            "asil_ay": txt_asil_ay.value.strip(),
            "asil_gun": txt_asil_gun.value.strip(),
            "suclar": [{"tarih": s["tarih_str"], "kasit": s["kasit"], "tur": s["tur"], "yil": str(s["toplam_gun"]//365), "ay": str((s["toplam_gun"]%365)//30), "gun": str((s["toplam_gun"]%365)%30)} for s in suclar_listesi]
        }
        try:
            with open("STGA_Mobil_Taslak.json", "w", encoding="utf-8") as f:
                json.dump(veri, f, ensure_ascii=False, indent=4)
            goster_mesaj("Başarılı", "Veriler 'STGA_Mobil_Taslak.json' dosyasına yedeklendi.", ft.Colors.PURPLE_400)
        except Exception as ex:
            goster_mesaj("Hata", f"Dışa aktarılamadı: {ex}", ft.Colors.RED_400)

    def taslak_ice_aktar(e):
        if not os.path.exists("STGA_Mobil_Taslak.json"):
            goster_mesaj("Yükleme Başarısız", "Bulunabilen kayıtlı 'STGA_Mobil_Taslak.json' dosyası yok!", ft.Colors.ORANGE_400)
            return
        try:
            with open("STGA_Mobil_Taslak.json", "r", encoding="utf-8") as f:
                veri = json.load(f)
            
            txt_asil_tarih.value = veri.get("asil_tarih", "")
            cmb_lehe_bitis.value = veri.get("lehe_bitis", "01.07.2016")
            chk_oran_var.value = (veri.get("oran_var", 0) == 1)
            txt_stt.value = veri.get("stt", "")
            txt_htt.value = veri.get("htt", "")
            txt_asil_yil.value = veri.get("asil_yil", "")
            txt_asil_ay.value = veri.get("asil_ay", "")
            txt_asil_gun.value = veri.get("asil_gun", "")
            
            suclar_listesi.clear()
            for s in veri.get("suclar", []):
                g_tarih = tarih_cevir(s["tarih"])
                t_gun = (rakam_al(s["yil"])*365) + (rakam_al(s["ay"])*30) + rakam_al(s["gun"])
                suclar_listesi.append({"tarih": g_tarih, "tarih_str": s["tarih"], "kasit": s["kasit"], "tur": s["tur"], "toplam_gun": t_gun})
            
            tabloyu_yenile()
            asil_ceza_kutularini_guncelle()
            goster_mesaj("Başarılı", "Taslak başarıyla geri yüklendi.", ft.Colors.GREEN_400)
        except Exception as ex:
            goster_mesaj("Hata", f"Yükleme hatası: {ex}", ft.Colors.RED_400)

    def rehber_goster(e):
        goster_mesaj("Sistem Rehberi", "1. Sol kısımdan asıl dosya bilgilerini girin.\n2. Dosyada 1/2 varsa işaretleyip sadece 1/2'lik cezaları girin.\n3. Denetimde işlenen yeni suçları ekleyip 'ŞTGA Hesapla' deyin.", ft.Colors.CYAN_400)

    def hakkında_goster(e):
        goster_mesaj("Hakkında", "Bu yazılım Recep BEGTAŞ (ab240611) ve Taha Yasin ÖNAL (ab256444) tarafından Diyarbakır İlamat ve İnfaz Bürosu için geliştirilmiştir.", ft.Colors.BLUE_400)

    def goster_oran_bilgi(e):
        goster_mesaj("İnfaz Oranı Hakkında", "Eğer elinizdeki müddetnamede hem 1/2 hem de 2/3 (veya 3/4) oranlı cezalar BİRLİKTE bulunuyorsa, bu kutuyu işaretleyiniz.", ft.Colors.CYAN_400)
    
    def goster_asil_bilgi(e):
        goster_mesaj("Asıl Ceza Alanı", "Bu alan varsayılan olarak KAPALIDIR. Yalnızca Asıl Suç Tarihi '01.06.2005' ile lehe yasa bitişi aralığında olduğunda aktif hale gelir. Sadece 1/2'lik cezaların toplamını giriniz.", ft.Colors.CYAN_400)

    # =========================================================================
    # GÖRSEL TASARIM ÖRGÜSÜ (KAPORTA) - "GÜVENLİ ATAMA" METODU UYGULANDI
    # =========================================================================
    
    # 1. Asıl Dosya Kutuları
    txt_asil_tarih = ft.TextField(label="Asıl Suç Tarihi", hint_text="GG.AA.YYYY", expand=True)
    txt_asil_tarih.on_blur = asil_tarih_degisti
    txt_asil_tarih.on_submit = asil_tarih_degisti
    
    cmb_lehe_bitis = ft.Dropdown(label="Lehe Yasa Bitişi", options=[ft.dropdown.Option("01.07.2016"), ft.dropdown.Option("30.03.2020")], value="01.07.2016", expand=True)
    cmb_lehe_bitis.on_change = asil_ceza_kutularini_guncelle
    
    chk_oran_var = ft.Checkbox(label="Dosyada 1/2 Oranı Var", value=False)
    chk_oran_var.on_change = chk_oran_degisti
    
    btn_oran_soru = ft.IconButton(icon=ft.Icons.HELP_OUTLINE, icon_color=ft.Colors.BLUE_400)
    btn_oran_soru.on_click = goster_oran_bilgi
    
    txt_stt = ft.TextField(label="Şartla Tahliye (ŞTT)", hint_text="GG.AA.YYYY", expand=True)
    txt_stt.on_blur = otomatik_tarih_formatla
    txt_stt.on_submit = otomatik_tarih_formatla
    
    txt_htt = ft.TextField(label="Bihakkın (HTT)", hint_text="GG.AA.YYYY", expand=True)
    txt_htt.on_blur = otomatik_tarih_formatla
    txt_htt.on_submit = otomatik_tarih_formatla
    
    txt_asil_yil = ft.TextField(label="Yıl", hint_text="0", disabled=True, expand=True)
    txt_asil_ay = ft.TextField(label="Ay", hint_text="0", disabled=True, expand=True)
    txt_asil_gun = ft.TextField(label="Gün", hint_text="0", disabled=True, expand=True)
    
    btn_asil_soru = ft.IconButton(icon=ft.Icons.HELP_OUTLINE, icon_color=ft.Colors.BLUE_400)
    btn_asil_soru.on_click = goster_asil_bilgi

    kart_asil_dosya = ft.Card(
        content=ft.Container(
            padding=12,
            content=ft.Column([
                ft.Text("📋 Bihakkın Dosyası Bilgileri", size=15, weight=ft.FontWeight.BOLD, color=ft.Colors.BLUE_200),
                ft.Row([txt_asil_tarih, cmb_lehe_bitis], spacing=10),
                ft.Row([chk_oran_var, btn_oran_soru], alignment=ft.MainAxisAlignment.SPACE_BETWEEN),
                ft.Row([txt_stt, txt_htt], spacing=10),
                ft.Row([txt_asil_yil, txt_asil_ay, txt_asil_gun, btn_asil_soru], spacing=5)
            ], spacing=8)
        )
    )

    # 2. Yeni Suç Kutuları
    txt_yeni_tarih = ft.TextField(label="Suç Tarihi", hint_text="GG.AA.YYYY", expand=True)
    txt_yeni_tarih.on_blur = otomatik_tarih_formatla
    txt_yeni_tarih.on_submit = otomatik_tarih_formatla
    
    cmb_kasit = ft.Dropdown(label="Kasıt?", options=[ft.dropdown.Option("Evet"), ft.dropdown.Option("Hayır")], value="Evet", expand=True)
    
    cmb_ceza_turu = ft.Dropdown(label="Tür", options=[ft.dropdown.Option("Hapis"), ft.dropdown.Option("APC"), ft.dropdown.Option("HAGB")], value="Hapis", expand=True)
    cmb_ceza_turu.on_change = hapis_sure_kutularini_guncelle
    
    txt_yeni_yil = ft.TextField(label="Yıl", hint_text="0", expand=True)
    txt_yeni_ay = ft.TextField(label="Ay", hint_text="0", expand=True)
    txt_yeni_gun = ft.TextField(label="Gün", hint_text="0", expand=True)
    
    btn_sucu_ekle = ft.ElevatedButton("Suçu Listeye Ekle", icon=ft.Icons.ADD, bgcolor=ft.Colors.GREEN_700, color=ft.Colors.WHITE, expand=True)
    btn_sucu_ekle.on_click = sucu_ekle_tetikle
    
    btn_sucu_sil = ft.ElevatedButton("Suçu Sil", icon=ft.Icons.DELETE, bgcolor=ft.Colors.RED_700, color=ft.Colors.WHITE, expand=True, disabled=True)
    btn_sucu_sil.on_click = sucu_sil_tetikle

    kart_yeni_suc = ft.Card(
        content=ft.Container(
            padding=12,
            content=ft.Column([
                ft.Text("🚨 Denetim Sürecinde İşlenen Yeni Suç", size=15, weight=ft.FontWeight.BOLD, color=ft.Colors.ORANGE_200),
                txt_yeni_tarih,
                ft.Row([cmb_kasit, cmb_ceza_turu], spacing=10),
                ft.Row([txt_yeni_yil, txt_yeni_ay, txt_yeni_gun], spacing=5),
                ft.Row([btn_sucu_ekle, btn_sucu_sil], spacing=10)
            ], spacing=8)
        )
    )

    tablo_suclar = ft.DataTable(
        columns=[
            ft.DataColumn(ft.Text("Tarih")),
            ft.DataColumn(ft.Text("Kasıt")),
            ft.DataColumn(ft.Text("Tür")),
            ft.DataColumn(ft.Text("Süre")),
        ],
        rows=[]
    )
    
    container_tablo = ft.Container(
        content=ft.Column([tablo_suclar], scroll=ft.ScrollMode.AUTO),
        height=130,
        border=ft.Border.all(1, ft.Colors.GREY_700),
        border_radius=8,
        padding=5
    )

    btn_hesapla = ft.ElevatedButton("⚖️ ŞTGA HESAPLA", bgcolor=ft.Colors.BLUE_700, color=ft.Colors.WHITE, height=45, expand=True)
    btn_hesapla.on_click = hesapla_tetikle
    
    btn_temizle = ft.ElevatedButton("🧹 Temizle", bgcolor=ft.Colors.RED_700, color=ft.Colors.WHITE, height=40, expand=True)
    btn_temizle.on_click = ekrani_temizle_tetikle
    
    btn_kaydet = ft.ElevatedButton("📄 Word Çıktısı", bgcolor=ft.Colors.GREEN_700, color=ft.Colors.WHITE, height=40, expand=True, disabled=True)
    btn_kaydet.on_click = evrak_secim_ekrani
    
    btn_disa_aktar = ft.IconButton(icon=ft.Icons.UPLOAD_FILE, tooltip="Dışa Aktar", icon_color=ft.Colors.PURPLE_300)
    btn_disa_aktar.on_click = taslak_disa_aktar
    
    btn_ice_aktar = ft.IconButton(icon=ft.Icons.DOWNLOAD_FOR_OFFLINE, tooltip="İçe Aktar", icon_color=ft.Colors.ORANGE_300)
    btn_ice_aktar.on_click = taslak_ice_aktar
    
    btn_nasil_calisir = ft.IconButton(icon=ft.Icons.LIGHTBULB_OUTLINE, tooltip="Nasıl Çalışır?", icon_color=ft.Colors.CYAN_300)
    btn_nasil_calisir.on_click = lambda e: goster_mesaj("Nasıl Çalışır?", "TCK 61, 7242 s.K. kıyaslamaları ve kronolojik tavan aşım kontrollerini otomatik yapar.")

    txt_sonuc = ft.TextField(
        label="Hesaplama Sonucu ve Hukuki Mütalaa", 
        multiline=True, 
        min_lines=12,  
        max_lines=25,  
        read_only=True, 
        text_style=ft.TextStyle(font_family="Consolas", size=13),
        border_color=ft.Colors.BLUE_400,
        expand=True  # Kutuya yatayda sonuna kadar genişleme yetkisi verdik!
    )

    lbl_rehber = ft.TextButton(
        "Sistem Rehberi", 
        on_click=rehber_goster,
        style=ft.ButtonStyle(color=ft.Colors.GREY_400)
    )
    
    lbl_ayrac = ft.Text(" | ", color=ft.Colors.GREY_600)
    
    lbl_hakkinda = ft.TextButton(
        "Hakkında", 
        on_click=hakkında_goster,
        style=ft.ButtonStyle(color=ft.Colors.GREY_400)
    )
    
    satir_linkler = ft.Row([lbl_rehber, lbl_ayrac, lbl_hakkinda], alignment=ft.MainAxisAlignment.END)

    # Ekran Elemanlarını Basma
    ana_kaydirici = ft.Column(scroll=ft.ScrollMode.AUTO, expand=True, spacing=12)
    ana_kaydirici.controls.extend([
        kart_asil_dosya,
        kart_yeni_suc,
        ft.Text("📋 Eklenen Suçlar Tablosu", size=13, weight=ft.FontWeight.BOLD, color=ft.Colors.GREY_400),
        container_tablo,
        ft.Row([btn_hesapla]),
        ft.Row([btn_temizle, btn_kaydet], spacing=10),
        ft.Row([btn_disa_aktar, btn_ice_aktar, btn_nasil_calisir], alignment=ft.MainAxisAlignment.CENTER, spacing=25),
        ft.Row([txt_sonuc]), # Kutuyu yatay şeride alarak ekranı tam kaplamasını sağladık
        satir_linkler
    ])

    page.add(ana_kaydirici)
    
    # Program açılır açılmaz gri kutu renklerini ayarlamak için ilk tetiklemeyi yapıyoruz
    asil_ceza_kutularini_guncelle()

ft.run(main)